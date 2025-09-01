# server.py
from __future__ import annotations

import os
import io
import re
import json as _json
from datetime import datetime, timedelta
from typing import List, Optional, Tuple, Dict

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from starlette.staticfiles import StaticFiles

# ================== Config ==================
BASE_URL = os.getenv("OPENAI_BASE_URL", "http://localhost:11434").rstrip("/")
DEFAULT_LLM = os.getenv("EF_AI_MODEL", "deepseek-r1:8b")
INDEX_DIR = os.getenv("INDEX_DIR", "faiss_index_EF")
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")
OUTPUT_DIR = os.getenv("OUTPUT_DIR", "outputs")
DESKTOP_XLSX = os.path.expanduser("~/Desktop/EF_PlanSemanal.xlsx")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ================== App ==================
app = FastAPI(title="EF-AI — Agents (Ollama local)")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)
templates = Jinja2Templates(directory="templates")
# Montar static solo si existe (evita RuntimeError)
if os.path.isdir("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

# ================== Dependencias ==================
import ollama
import pypdf, docx2txt
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.docstore.document import Document

# Excel
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

# Embeddings MULTILINGÜE
_EMBEDDINGS = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")

THINK_OPEN, THINK_CLOSE = "<think>", "</think>"

def strip_think_streaming():
    """Crea un filtro que elimina en streaming cualquier bloque <think>...</think>."""
    state = {"in": False, "saw": False}
    def _f(tok: str) -> str:
        t = tok
        out = ""
        while t:
            if not state["in"]:
                i = t.find(THINK_OPEN)
                if i == -1:
                    out += t
                    t = ""
                else:
                    out += t[:i]
                    t = t[i + len(THINK_OPEN):]
                    state["in"] = True
                    state["saw"] = True
            else:
                j = t.find(THINK_CLOSE)
                if j == -1:
                    # seguimos dentro de <think> — descarta resto
                    return out
                t = t[j + len(THINK_CLOSE):]
                state["in"] = False
        return out
    return _f, state

def scrub_pii(x: str) -> str:
    return x.strip()

# ---------- FAISS index ----------
def _split_text(text: str) -> List[str]:
    return RecursiveCharacterTextSplitter(
        chunk_size=900, chunk_overlap=150, separators=["\n\n", "\n", " ", ""]
    ).split_text(text)

def _read_upload_file(f: UploadFile) -> Tuple[str, bytes, str]:
    """Lee bytes y extrae texto. Devuelve (nombre, raw_bytes, texto)."""
    name = f.filename or "archivo"
    raw = f.file.read()
    text = ""
    if name.lower().endswith(".pdf"):
        reader = pypdf.PdfReader(io.BytesIO(raw))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif name.lower().endswith((".docx", ".doc")):
        tmp = os.path.join(UPLOAD_DIR, f"__tmp_{os.getpid()}_{name}")
        with open(tmp, "wb") as w:
            w.write(raw)
        try:
            text = docx2txt.process(tmp) or ""
        finally:
            try:
                os.remove(tmp)
            except Exception:
                pass
    else:
        try:
            text = raw.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    # reset pointer
    try:
        f.file.seek(0)
    except Exception:
        pass
    return name, raw, text

def load_index() -> FAISS:
    if not os.path.isdir(INDEX_DIR):
        raise RuntimeError("INDEX_NOT_FOUND")
    return FAISS.load_local(INDEX_DIR, _EMBEDDINGS, allow_dangerous_deserialization=True)

def save_index(db: FAISS) -> None:
    db.save_local(INDEX_DIR)

def add_to_index(files: List[UploadFile], unidad: str = "General") -> int:
    new_docs = []
    for f in files:
        if not (f and f.filename):
            continue
        name, raw, text = _read_upload_file(f)
        # guarda una copia EXACTA en uploads
        try:
            with open(os.path.join(UPLOAD_DIR, name), "wb") as w:
                w.write(raw)
        except Exception:
            pass
        for chunk in _split_text(text):
            new_docs.append(Document(page_content=chunk, metadata={"source": os.path.join(UPLOAD_DIR, name), "unidad": unidad}))
    if not new_docs:
        return 0
    if os.path.isdir(INDEX_DIR):
        db = load_index()
        db.add_documents(new_docs)
    else:
        db = FAISS.from_documents(new_docs, _EMBEDDINGS)
    save_index(db)
    return len(new_docs)

def _retrieve_with_fallback(db: FAISS, q: str, k: int, unidad: Optional[str]):
    res = db.similarity_search_with_score(q, k=24)
    def pick(rs, unit):
        out = []
        for d, sc in rs:
            if unit and d.metadata.get("unidad") != unit:
                continue
            out.append((d, sc))
            if len(out) == k:
                break
        return out
    got = pick(res, unidad)
    if not got:
        got = pick(res, None)
    return got

def _build_sources(pairs):
    items = []
    for (doc, score) in pairs:
        txt = (doc.page_content or "").replace("\n", " ")
        snippet = (txt[:240] + "…") if len(txt) > 240 else txt
        items.append({
            "source": doc.metadata.get("source", "?"),
            "unidad": doc.metadata.get("unidad", "?"),
            "score": float(score),
            "snippet": snippet
        })
    return items

def _is_context_weak(context: str) -> bool:
    letters = sum(c.isalpha() for c in context)
    return letters < 120

# ---------- Excel & Planificador ----------
DAYS_MAP = {
    "lunes":0,"martes":1,"miércoles":2,"miercoles":2,"jueves":3,"viernes":4,"sábado":5,"sabado":5,"domingo":6,
    "mon":0,"tue":1,"wed":2,"thu":3,"fri":4,"sat":5,"sun":6
}
DAY_NAMES = ["Lunes","Martes","Miércoles","Jueves","Viernes","Sábado","Domingo"]

COURSE_PATTERNS = [
    (r"\b(1|1º|primero)\s*eso\b", "1ESO"),
    (r"\b(2|2º|segundo)\s*eso\b", "2ESO"),
    (r"\b(3|3º|tercero)\s*eso\b", "3ESO"),
    (r"\b(4|4º|cuarto)\s*eso\b", "4ESO"),
    (r"\bbach(illerato)?\b", "BACHI"),
    (r"\bprimaria\b", "PRIMARIA"),
]

def detect_course(text: str, default_course: str) -> str:
    t = (text or "").lower()
    for pat, lab in COURSE_PATTERNS:
        if re.search(pat, t):
            return lab
    return default_course

def extract_sessions_numbered_with_desc(text: str) -> List[Tuple[int, str, str]]:
    """Devuelve [(num, titulo, desc)] ordenado por num usando cabeceras 'Sx — Título'."""
    text = text or ""
    header_re = re.compile(r"^\s*S\s*([0-9]{1,2})\s*[—–\-]\s*(.+)$", re.MULTILINE)
    matches = list(header_re.finditer(text))
    out: List[Tuple[int,str,str]] = []
    for i, m in enumerate(matches):
        n = int(m.group(1))
        title = m.group(2).strip()
        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        desc = text[start:end].strip()
        desc = re.sub(r"\s+", " ", desc)
        if len(desc) > 500:
            desc = desc[:500].rstrip() + "…"
        if title:
            out.append((n, title, desc))
    out.sort(key=lambda x: x[0])
    return out

def extract_sessions_simple(text: str, filename: str) -> List[str]:
    titles = []
    for line in (text or "").splitlines():
        ln = line.strip()
        if re.search(r"\bsesi(ó|o)n\b|\bsession\b", ln, re.IGNORECASE):
            ln = re.sub(r"^[\-\*\d\.\)]\s*", "", ln)
            titles.append(ln)
    if not titles:
        base = os.path.splitext(os.path.basename(filename))[0]
        titles = [base]
    out = []
    for t in titles:
        t = re.sub(r"\s{2,}", " ", t).strip()
        if t and (not out or t != out[-1]):
            out.append(t)
    return out

def parse_time_ranges(times_csv: str) -> List[str]:
    parts = [p.strip() for p in re.split(r"[,\;]", times_csv or "") if p.strip()]
    uniq = list(dict.fromkeys(parts))
    out = []
    for p in uniq:
        if not re.match(r"^\d{2}:\d{2}\s*-\s*\d{2}:\d{2}$", p):
            raise ValueError(f"Horario inválido: '{p}' (HH:MM-HH:MM)")
        out.append(p.replace(" ", ""))
    if not out:
        out = ["08:00-09:00"]
    return out

def monday_of_week(d: datetime) -> datetime:
    return d - timedelta(days=d.weekday())

_PALETTE = ["DDEBF7","E2F0D9","FFF2CC","F8CBAD","E4DFEC","D9D2E9","FCE4D6","EAD1DC","C9DAF8","D0E0E3"]
def _color_for(course: str) -> str:
    return _PALETTE[abs(hash(course)) % len(_PALETTE)]

def ensure_week_sheet(wb: Workbook, course: str, week_num: int,
                      week_dates: Dict[int, str], time_ranges: List[str]):
    title = f"{course} — Semana {week_num}"
    if title in wb.sheetnames:
        wb.remove(wb[title])
    ws = wb.create_sheet(title)

    headers = ["Hora", "Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]
    dates_row = ["", week_dates.get(0,""), week_dates.get(1,""), week_dates.get(2,""),
                    week_dates.get(3,""), week_dates.get(4,"")]

    head_font = Font(bold=True)
    head_fill = PatternFill("solid", fgColor="BDD7EE")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin = Side(style="thin", color="BBBBBB")
    border = Border(left=thin,right=thin,top=thin,bottom=thin)

    ws.append(headers); ws.append(dates_row)

    ws.column_dimensions["A"].width = 16
    for col in "BCDEF": ws.column_dimensions[col].width = 34

    for col_idx in range(1, 7):
        c1 = ws.cell(row=1, column=col_idx)
        c2 = ws.cell(row=2, column=col_idx)
        c1.font=head_font; c1.fill=head_fill; c1.alignment=center; c1.border=border
        c2.font=head_font; c2.alignment=center; c2.border=border

    r = 3
    for tr in time_ranges:
        ws.cell(row=r, column=1, value=tr).alignment=center
        ws.cell(row=r, column=1).border=border
        for col_idx in range(2, 7):
            cell = ws.cell(row=r, column=col_idx)
            cell.alignment=center
            cell.border=border
        r += 1
    return ws

def write_title_desc(cell, title: str, desc: str):
    clean_desc = (desc or "").strip()
    try:
        # openpyxl >= 3.1 con RichText
        from openpyxl.cell.rich_text import CellRichText, TextBlock, InlineFont
        rt = CellRichText()
        rt.add(TextBlock(InlineFont(b=True), title))
        if clean_desc:
            rt.add(TextBlock(InlineFont(b=False), "\n" + clean_desc))
        cell.value = rt
    except Exception:
        # fallback: texto con salto de línea y font bold para todo
        cell.value = title + ("\n" + clean_desc if clean_desc else "")
        cell.font = Font(bold=True)

def _parse_week_num(name: str, course: str) -> Optional[int]:
    m = re.match(rf"^{re.escape(course)}\s*—\s*Semana\s*(\d+)$", name)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    return None

def find_last_week_info(wb, course: str) -> Tuple[int, Optional[datetime]]:
    max_n = 0
    last_monday = None
    for name in wb.sheetnames:
        n = _parse_week_num(name, course)
        if n and n > max_n:
            max_n = n
            ws = wb[name]
            try:
                date_str = str(ws.cell(row=2, column=2).value or "")
                if re.match(r"^\d{4}-\d{2}-\d{2}$", date_str):
                    last_monday = datetime.strptime(date_str, "%Y-%m-%d")
            except Exception:
                pass
    return max_n, last_monday

def plan_sessions_weekly_sheets(
    course: str,
    session_items: List[Tuple[str, str]],   # [(title, desc)]
    start_date: str,
    days_csv: str,
    times_csv: str,
    continue_from_last: bool = False
) -> Dict[str,int]:

    if os.path.exists(DESKTOP_XLSX):
        wb = load_workbook(DESKTOP_XLSX)
    else:
        wb = Workbook()
        if wb.active.title == "Sheet":
            wb.active.title = "Resumen"

    base_dt = datetime.strptime(start_date, "%Y-%m-%d")
    days_tokens = [t.strip().lower() for t in re.split(r"[,\;]", days_csv) if t.strip()]
    day_idx_sorted = sorted({DAYS_MAP[t] for t in days_tokens if t in DAYS_MAP})
    if len(day_idx_sorted) == 0:
        day_idx_sorted = [0, 2]  # Lunes/Miércoles
    if len(day_idx_sorted) == 1:
        day_idx_sorted = [day_idx_sorted[0], day_idx_sorted[0]]

    time_ranges = parse_time_ranges(times_csv)
    same_row = (len(time_ranges) == 1) or (len(time_ranges) >= 2 and time_ranges[0] == time_ranges[1])

    if continue_from_last:
        last_n, last_mon = find_last_week_info(wb, course)
        week_num = last_n + 1
        week_num_start = week_num
        m0 = (last_mon + timedelta(days=7)) if last_mon else monday_of_week(base_dt)
    else:
        to_delete = [s for s in wb.sheetnames if s.startswith(f"{course} — Semana ")]
        for s in to_delete:
            wb.remove(wb[s])
        week_num = 1
        week_num_start = 1
        m0 = monday_of_week(base_dt)

    added = 0
    i = 0
    while i < len(session_items):
        week_start = m0 + timedelta(weeks=(week_num - week_num_start))
        week_dates = {d: (week_start + timedelta(days=d)).strftime("%Y-%m-%d") for d in range(5)}

        ws = ensure_week_sheet(wb, course, week_num, week_dates, time_ranges)

        first_row = 3
        second_row = 3 if same_row else 4

        # Sesión 1
        if i < len(session_items):
            title, desc = session_items[i]
            col = 2 + day_idx_sorted[0]
            cell = ws.cell(row=first_row, column=col)
            write_title_desc(cell, f"S{i+1}: {title}", desc)
            cell.fill = PatternFill("solid", fgColor=_color_for(course))
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            added += 1; i += 1

        # Sesión 2
        if i < len(session_items):
            title, desc = session_items[i]
            col = 2 + day_idx_sorted[1]
            cell = ws.cell(row=second_row, column=col)
            write_title_desc(cell, f"S{i+1}: {title}", desc)
            cell.fill = PatternFill("solid", fgColor=_color_for(course))
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            added += 1; i += 1

        week_num += 1

    wb.save(DESKTOP_XLSX)
    try:
        wb.save(os.path.join(OUTPUT_DIR, "EF_PlanSemanal.xlsx"))
    except Exception:
        pass

    return {
        "semana_inicial": week_num_start,
        "semanas_creadas": (week_num - week_num_start),
        "sesiones_colocadas": added
    }

# ================== Rutas ==================
@app.get("/")
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "default_model": DEFAULT_LLM})

@app.get("/api/health")
async def api_health():
    try:
        client = ollama.Client(host=BASE_URL)
        models_raw = client.list().get("models", [])
        # normaliza a lista de nombres
        names = []
        for m in models_raw:
            name = m.get("model") or m.get("name") or str(m)
            names.append(name)
        return {"status":"ok","models":names,"base_url":BASE_URL}
    except Exception as e:
        return {"status":"down","error":str(e),"base_url":BASE_URL}

@app.get("/files/{path:path}")
async def serve_file(path: str):
    full = os.path.join(OUTPUT_DIR, os.path.basename(path))
    if os.path.exists(full):
        return FileResponse(full)
    # también mirar escritorio copia si existe
    desk = os.path.expanduser("~/Desktop/" + os.path.basename(path))
    if os.path.exists(desk):
        return FileResponse(desk)
    return JSONResponse({"ok": False, "error":"FILE_NOT_FOUND"}, status_code=404)

@app.post("/api/index")
async def api_index(unit: str = Form("General"), files: List[UploadFile] = File(...)):
    try:
        added = add_to_index(files, unidad=(unit or "General"))
        return {"ok": True, "added": added, "unit": unit}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# ---- Chat RAG (stream) ----
@app.post("/api/ask_stream")
async def api_ask_stream(payload: dict):
    question = (payload.get("question") or "").strip()
    unit = payload.get("unit") or None
    k = int(payload.get("k", 3))
    model = payload.get("model") or DEFAULT_LLM
    fallback = bool(payload.get("fallback", False))
    if not question:
        return JSONResponse({"ok": False, "error": "EMPTY_QUESTION"}, status_code=400)
    q = scrub_pii(question)

    # intenta cargar índice
    try:
        db = load_index()
    except Exception:
        db = None

    if db is None:
        if not fallback:
            def _e():
                yield f"data: {_json.dumps({'final':'No está en el contexto','sources':[]})}\n\n"
                yield "event: end\ndata: {}\n\n"
            return StreamingResponse(_e(), media_type="text/event-stream")
        # sin índice → respuesta general
        def _gen():
            client = ollama.Client(host=BASE_URL)
            filt, state = strip_think_streaming()
            acc = ""
            for ch in client.chat(
                model=model, stream=True,
                messages=[{"role":"system","content":"Eres profesor de Educación Física. Responde breve y claro en español."},
                          {"role":"user","content":q}],
                options={"temperature":0.2,"num_ctx":4096,"num_predict":512}
            ):
                tok = ch.get("message",{}).get("content","")
                vis = filt(tok)
                if vis:
                    acc += vis
                    yield f"data: {_json.dumps({'token': vis})}\n\n"
            yield f"data: {_json.dumps({'final': acc.strip(), 'sources': [], 'saw_think': state['saw']})}\n\n"
            yield "event: end\ndata: {}\n\n"
        return StreamingResponse(_gen(), media_type="text/event-stream")

    # recupera contexto
    pairs = _retrieve_with_fallback(db, q, k, unit)
    sources = _build_sources(pairs)

    if not pairs and not fallback:
        def _empty():
            yield f"data: {_json.dumps({'final':'No está en el contexto','sources':sources})}\n\n"
            yield "event: end\ndata: {}\n\n"
        return StreamingResponse(_empty(), media_type="text/event-stream")

    context = "\n\n".join(d.page_content for (d, _) in pairs) if pairs else ""
    weak = _is_context_weak(context)

    # si el contexto es débil y se permite fallback, responde general
    if weak and fallback:
        def _gen():
            client = ollama.Client(host=BASE_URL)
            filt, state = strip_think_streaming()
            acc = ""
            for ch in client.chat(
                model=model, stream=True,
                messages=[{"role":"system","content":"Eres profesor de Educación Física. Responde breve y claro en español."},
                          {"role":"user","content":q}],
                options={"temperature":0.2,"num_ctx":4096,"num_predict":512}
            ):
                tok = ch.get("message",{}).get("content","")
                vis = filt(tok)
                if vis:
                    acc += vis
                    yield f"data: {_json.dumps({'token': vis})}\n\n"
            yield f"data: {_json.dumps({'final': acc.strip(), 'sources': sources, 'saw_think': state['saw']})}\n\n"
            yield "event: end\ndata: {}\n\n"
        return StreamingResponse(_gen(), media_type="text/event-stream")

    # respuesta basada EN CONTEXTO
    sys_prompt = (
        "Eres profesor de Educación Física. Responde SOLO con la información del CONTEXTO. "
        "Si no hay información relevante, responde exactamente: 'No está en el contexto'. "
        "NO muestres razonamiento ni etiquetas 'think'. Respuesta final breve en español."
    )
    user_prompt = f"<CONTEXTO>\n{context}\n</CONTEXTO>\n<PREGUNTA>\n{q}\n</PREGUNTA>"

    def _gen_ctx():
        client = ollama.Client(host=BASE_URL)
        filt, state = strip_think_streaming()
        acc = ""
        for ch in client.chat(
            model=model, stream=True,
            messages=[{"role":"system","content":sys_prompt},
                      {"role":"user","content":user_prompt}],
            options={"temperature":0.15,"num_ctx":4096,"num_predict":512}
        ):
            tok = ch.get("message",{}).get("content","")
            vis = filt(tok)
            if vis:
                acc += vis
                yield f"data: {_json.dumps({'token': vis})}\n\n"
        yield f"data: {_json.dumps({'final': acc.strip(),'sources': sources,'saw_think': state['saw']})}\n\n"
        yield "event: end\ndata: {}\n\n"

    return StreamingResponse(_gen_ctx(), media_type="text/event-stream")

# --- Rúbrica / Corrector simple endpoints ---
@app.post("/api/rubric")
async def api_rubric(yaml_text: str = Form(""), course: str = Form(""), unit: str = Form("")):
    from docx import Document
    doc = Document()
    doc.add_heading(f"Rúbrica — {course} — {unit}", level=1)
    doc.add_paragraph(yaml_text or "(YAML no proporcionado)")
    out_path = os.path.join(OUTPUT_DIR, f"rubrica_{course}_{unit}.docx".replace(" ", "_"))
    doc.save(out_path)
    return {"ok": True, "path": f"/files/{os.path.basename(out_path)}"}

@app.post("/api/correct")
async def api_correct(model: str = Form(DEFAULT_LLM),
                      rubric_text: str = Form(""),
                      student_text: str = Form("")):
    client = ollama.Client(host=BASE_URL)
    sys = "Eres profesor de Educación Física. Evalúa el trabajo según la rúbrica dada y ofrece feedback claro y breve en español."
    user = f"RÚBRICA:\n{rubric_text}\n\nTRABAJO DEL ESTUDIANTE:\n{student_text}"
    try:
        resp = client.chat(
            model=model,
            messages=[{"role":"system","content":sys},
                      {"role":"user","content":user}],
            options={"temperature":0.2,"num_ctx":4096,"num_predict":512}
        )
        txt = resp.get("message",{}).get("content","").strip()
        return {"ok": True, "result": txt}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# --- Planificador: endpoint ---
@app.post("/api/planify")
async def api_planify(
    course: str = Form(...),             # curso base por defecto
    start_date: str = Form(...),         # YYYY-MM-DD
    days_csv: str = Form(...),           # "Lunes, Miércoles"
    times_csv: str = Form(...),          # "08:00-09:00, 12:25-13:20"
    use_llm: bool = Form(False),
    continue_from_last: bool = Form(False),
    files: List[UploadFile] = File(...),
):
    try:
        sessions_by_course: Dict[str, List[Tuple[str,str]]] = {}
        for f in files:
            name, raw, text = _read_upload_file(f)
            detected_course = detect_course(text, course)
            numbered = extract_sessions_numbered_with_desc(text)
            if numbered:
                items = [(title, desc) for (_n, title, desc) in numbered]
            else:
                items = [(t, "") for t in extract_sessions_simple(text, name)]
            if items:
                sessions_by_course.setdefault(detected_course, []).extend(items)

        if not sessions_by_course:
            return {"ok": False, "error": "No se detectaron sesiones en los archivos."}

        resumen = {}
        for curso, items in sessions_by_course.items():
            res = plan_sessions_weekly_sheets(
                course=curso,
                session_items=items,
                start_date=start_date,
                days_csv=days_csv,
                times_csv=times_csv,
                continue_from_last=bool(continue_from_last)
            )
            resumen[curso] = res

        return {
            "ok": True,
            "desktop_path": DESKTOP_XLSX,
            "download": "/files/EF_PlanSemanal.xlsx",
            "resumen": resumen
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}

# --- Agente autónomo ---
from ef_agent import run_agent  # asegúrate de tener ef_agent.py en la misma carpeta

@app.post("/api/agent")
async def api_agent(goal: str = Form(...), model: str = Form(DEFAULT_LLM)):
    try:
        result = run_agent(goal=goal, model=model, max_steps=12)
        return result
    except Exception as e:
        return {"ok": False, "error": str(e)}