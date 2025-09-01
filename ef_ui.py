# -*- coding: utf-8 -*-
"""
EF-AI — Multi-agent UI (Gradio) sin Guardrails
Pestañas:
  - Chat RAG (consulta índice)
  - Indexador (subir/añadir docs)
  - Generador de Rúbricas (YAML -> DOCX)
  - Corrector (aplica rúbrica sobre texto del alumno)
  - Ajustes (modelos Ollama, salud del servidor)
"""

import os, shutil, json, textwrap
from pathlib import Path
from datetime import datetime
from typing import List, Optional

import requests
import gradio as gr
from dotenv import load_dotenv

# Carga variables .env
load_dotenv()

# --- RAG backend ---
from ef_rag import INDEX_DIR, add_to_index, load_index, make_qa, DEFAULT_MODEL

# PII scrub opcional
try:
    from privacy import scrub as scrub_pii
except Exception:
    def scrub_pii(x: str) -> str:
        return x

# --------- Constantes / rutas ----------
UPLOAD_STAGING = Path("uploads")
UPLOAD_STAGING.mkdir(exist_ok=True)
OUTPUTS_DIR = Path("outputs")
OUTPUTS_DIR.mkdir(exist_ok=True)

BASE_URL = os.getenv("OPENAI_BASE_URL", "http://localhost:11434")
DEFAULT_LLM = os.getenv("EF_AI_MODEL", DEFAULT_MODEL)

# --------- Utilidades Ollama ----------
def ollama_models() -> List[str]:
    """Devuelve la lista de modelos instalados en Ollama (o fallback)."""
    try:
        r = requests.get(f"{BASE_URL}/api/tags", timeout=2)
        r.raise_for_status()
        data = r.json()
        models = [m["name"] for m in data.get("models", [])]
        return sorted(set(models)) or [DEFAULT_LLM]
    except Exception:
        return [DEFAULT_LLM]

def ollama_health() -> str:
    try:
        r = requests.get(f"{BASE_URL}/api/tags", timeout=2)
        r.raise_for_status()
        return "🟢 Ollama: OK"
    except requests.exceptions.ConnectionError:
        return "🔴 Ollama: servidor no accesible (arranca con: `ollama serve`)"
    except Exception as e:
        return f"🟠 Ollama: {type(e).__name__}"

# --------- Indexado ----------
def index_uploads(files, unidad):
    if not files:
        return gr.Info("Sube algún archivo primero."), None
    saved = []
    for f in files:
        dst = UPLOAD_STAGING / f.name
        shutil.copyfile(f.name, dst)
        saved.append(str(dst))
    meta = {"unidad": unidad or "General"}
    add_to_index(saved, default_meta=meta)
    return gr.Success(f"Indexados {len(saved)} archivo(s) en '{INDEX_DIR}'."), None

def clear_session_index():
    for p in UPLOAD_STAGING.glob("*"):
        try:
            p.unlink()
        except Exception:
            pass
    return gr.Warning("Subidas temporales limpiadas (el índice FAISS global no se borra).")

# --------- Chat RAG ----------
def ask_rag(history, question, unidad, k, model_name):
    q = (question or "").strip()
    if not q:
        return history, gr.Info("Escribe una pregunta."), f"Usando k={k}. Contextos: 0.", ""
    q = scrub_pii(q)

    # Carga índice
    try:
        db = load_index()
    except Exception as e:
        return history, gr.Error("No existe el índice. Añade documentos primero."), f"Error: {e}", ""

    # QA
    model = model_name or DEFAULT_LLM
    qa = make_qa(db, model_name=model, temperature=0.2)
    try:
        ans = qa.run(q, k=int(k or 3), unidad=(unidad or None), show=False)
    except Exception as e:
        return history, gr.Error(f"Error LLM: {e}"), f"Modelo: {model}", ""

    # Chatbot (messages)
    history = history + [
        {"role": "user", "content": q},
        {"role": "assistant", "content": ans},
    ]
    context_text = f"Usando k={k}. Modelo: {model}"
    return history, None, context_text, ""

# --------- Generador de rúbricas ----------
def generate_rubric_from_yaml(yaml_text: str, course: str, unit: str) -> str:
    """Crea un DOCX simple a partir de YAML de rúbrica."""
    import yaml
    from docx import Document

    if not yaml_text.strip():
        raise ValueError("Pegue un YAML de rúbrica.")
    y = yaml.safe_load(yaml_text)

    doc = Document()
    doc.add_heading(f"Rubric – {course} – {unit}", level=1)

    dims = y.get("dimensions", [])
    if not isinstance(dims, list):
        dims = []

    for block in dims:
        name = block.get("name", "Dimension")
        doc.add_heading(name, level=2)
        for lvl in block.get("levels", []):
            label = lvl.get("label", "Level")
            desc = lvl.get("desc", "")
            doc.add_paragraph(f"{label}: {desc}")

    OUTPUTS_DIR.mkdir(exist_ok=True)
    out = OUTPUTS_DIR / f"rubric_{course}_{unit}_{datetime.now():%Y%m%d_%H%M}.docx"
    doc.save(out)
    return str(out)

# --------- Corrector (rúbrica + tarea) ----------
def correct_one_llm(model_name: str, rubric_text: str, student_text: str) -> str:
    """Invoca el LLM con un prompt compacto para devolver JSON con score & reasons."""
    from langchain_ollama import ChatOllama
    if not rubric_text.strip() or not student_text.strip():
        return json.dumps({"error": "Faltan rúbrica o texto del estudiante."}, ensure_ascii=False)

    sys_prompt = (
        "Eres un evaluador de Educación Física en Secundaria. Evalúa la tarea usando esta rúbrica (escala 1–4).\n"
        "Devuelve SOLO JSON con las claves: score (1-4) y reasons (≤40 palabras, español, EF)."
    )
    user = (
        f"RÚBRICA (YAML):\n{rubric_text}\n\n"
        f"TAREA DEL ESTUDIANTE:\n{student_text}\n\n"
        "Estructura JSON exacta:\n"
        "{\"score\": <1-4>, \"reasons\": \"<≤40 palabras>\"}"
    )
    llm = ChatOllama(model=model_name or DEFAULT_LLM, temperature=0.2)
    msg = f"<sistema>\n{sys_prompt}\n</sistema>\n<usuario>\n{user}\n</usuario>"
    try:
        out = llm.invoke(msg).content.strip()
    except Exception as e:
        out = json.dumps({"error": f"LLM: {e}"}, ensure_ascii=False)
    return out

# --------- UI  ----------
THEME_CSS = """
:root {
  --glass-bg: rgba(255,255,255,0.7);
  --glass-brd: rgba(255,255,255,0.35);
  --accent: #ff7a59;
}
.gradio-container {background: linear-gradient(135deg, #eef2ff 0%, #e0f7ff 100%) fixed;}
.card {
  background: var(--glass-bg);
  backdrop-filter: blur(10px);
  border: 1px solid var(--glass-brd);
  border-radius: 14px;
  padding: 14px;
}
#titlebar {font-weight:700; font-size: 22px;}
.btn-primary {background: var(--accent) !important; border: none;}
"""

def build_ui():
    with gr.Blocks(theme=gr.themes.Soft(), css=THEME_CSS, title="EF-AI — Agents") as demo:
        # Header con salud
        with gr.Row():
            with gr.Column():
                gr.Markdown("<div id='titlebar'>🏫 EF-AI — Agents (Ollama local)</div>")
                health = gr.Markdown(ollama_health())
        with gr.Tabs():
            # ---------------- Chat RAG ----------------
            with gr.Tab("🔎 Chat RAG"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### Controles", elem_classes=["card"])
                        unidad = gr.Textbox(value="General", label="Unidad/Tag")
                        k = gr.Slider(1, 6, value=3, step=1, label="k")
                        model_choice = gr.Dropdown(
                            choices=ollama_models(),
                            value=DEFAULT_LLM,
                            label="Modelo Ollama"
                        )
                    with gr.Column(scale=2):
                        chat = gr.Chatbot(type="messages", height=420, elem_classes=["card"])
                        with gr.Row():
                            question = gr.Textbox(placeholder="Escribe tu query...", label="Pregunta")
                            ask_btn = gr.Button("Preguntar", variant="primary", elem_classes=["btn-primary"])
                        footer = gr.Markdown("Usando k=3. Modelo: "+DEFAULT_LLM)
                        debug_box = gr.Textbox(label="Logs breves", interactive=False)

                ask_btn.click(
                    ask_rag,
                    [chat, question, unidad, k, model_choice],
                    [chat, health, footer, debug_box]
                )

            # ---------------- Indexador ----------------
            with gr.Tab("📚 Indexador"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### Añadir documentos al índice", elem_classes=["card"])
                        unidad_idx = gr.Textbox(value="General", label="Unidad/Tag")
                        files = gr.Files(label="Adjunta PDF/DOCX/TXT/MD", file_count="multiple")
                        add_btn = gr.Button("➕ Añadir al índice", variant="primary", elem_classes=["btn-primary"])
                        clear_btn = gr.Button("🧹 Limpiar subidas (no borra FAISS)")
                        status_idx = gr.Markdown()
                add_btn.click(index_uploads, [files, unidad_idx], [status_idx, files])
                clear_btn.click(clear_session_index, outputs=[status_idx])

            # ---------------- Rúbricas ----------------
            with gr.Tab("🧩 Generador de rúbricas"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### Config", elem_classes=["card"])
                        course = gr.Textbox(value="4º ESO", label="Curso")
                        unit = gr.Textbox(value="Futsal", label="Unidad")
                    with gr.Column(scale=2):
                        yaml_box = gr.Code(label="Pega aquí el YAML de la rúbrica", language="yaml", lines=20)
                        gen_btn = gr.Button("Generar DOCX", variant="primary", elem_classes=["btn-primary"])
                        out_path = gr.Textbox(label="Archivo generado", interactive=False)
                def _gen(ytext, c, u):
                    try:
                        p = generate_rubric_from_yaml(ytext, c, u)
                        return gr.Success("Rúbrica generada."), p
                    except Exception as e:
                        return gr.Error(str(e)), ""
                gen_btn.click(_gen, [yaml_box, course, unit], [health, out_path])

            # ---------------- Corrector ----------------
            with gr.Tab("✅ Corrector (rúbrica + tarea)"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### Modelo", elem_classes=["card"])
                        model_corr = gr.Dropdown(choices=ollama_models(), value=DEFAULT_LLM, label="Modelo Ollama")
                    with gr.Column(scale=2):
                        rubric_text = gr.Code(label="Rúbrica (YAML)", language="yaml", lines=16)
                        student_text = gr.Textbox(label="Tarea del estudiante", lines=6, placeholder="Pega aquí la respuesta del alumno...")
                        corr_btn = gr.Button("Corregir", variant="primary", elem_classes=["btn-primary"])
                        corr_out = gr.Code(label="Resultado JSON", language="json", lines=10)
                corr_btn.click(lambda m, r, s: correct_one_llm(m, r, s),
                               [model_corr, rubric_text, student_text],
                               [corr_out])

            # ---------------- Ajustes ----------------
            with gr.Tab("⚙️ Ajustes"):
                with gr.Row():
                    with gr.Column(scale=1):
                        gr.Markdown("#### Estado del servidor", elem_classes=["card"])
                        refresh_btn = gr.Button("🔄 Comprobar Ollama")
                        models_box = gr.JSON(label="Modelos instalados")
                    with gr.Column(scale=1):
                        gr.Markdown("#### Variables activas", elem_classes=["card"])
                        env_box = gr.JSON(value={
                            "OPENAI_BASE_URL": os.getenv("OPENAI_BASE_URL", ""),
                            "EF_AI_MODEL": DEFAULT_LLM,
                            "INDEX_DIR": str(INDEX_DIR),
                        })

                def _refresh():
                    return ollama_health(), ollama_models()
                refresh_btn.click(_refresh, None, [health, models_box])

    return demo


if __name__ == "__main__":
    demo = build_ui()
    demo.launch()